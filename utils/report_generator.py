import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.logger import get_logger

logger = get_logger("ReportGenerator")

# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, italic=False, color: str = None,
               size: int = 10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)


def _heading(doc, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(6)
    h.paragraph_format.space_after = Pt(4)
    return h


def _add_image(doc, path: str, width_inches: float = 5.8, caption: str = ""):
    if path and os.path.exists(path):
        doc.add_picture(path, width=Inches(width_inches))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cap.runs:
                run.font.italic = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    else:
        doc.add_paragraph("⚠ Screenshot not available.")


def _status_banner(doc, passed: bool):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    if passed:
        _set_cell_bg(cell, "375623")
        _cell_text(cell, "  ✅  STATUS: PASSED", bold=True, color="FFFFFF",
                   size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        _set_cell_bg(cell, "C00000")
        _cell_text(cell, "  ❌  STATUS: FAILED", bold=True, color="FFFFFF",
                   size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def _summary_table(doc, rows: list[tuple]):
    """Render a 2-column (label | value) summary table."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        lc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        _set_cell_bg(lc, "2F5496")
        _cell_text(lc, label, bold=True, color="FFFFFF", size=10)
        lc.width = Inches(2.0)
        _cell_text(vc, value, size=10)
    doc.add_paragraph()


# ── Step table ────────────────────────────────────────────────────────────────

STEP_EXPECTED = {
    1:  "Dashboard page loads successfully after login",
    2:  "ParQ Your Data page visible with Data Lakehouse section and Create New button",
    3:  "Choose Source page loads with Cloud Storage, Database & Warehouse, Governance, API tabs",
    4:  "Governance upload area ('Click to Upload') is visible",
    5:  "File uploads successfully and filename appears in the UI",
    6:  "Next button is enabled after upload and navigates to File Setup page without errors",
    7:  "Metadata tab shows a populated table with at least one row",
    8:  "Sample Data tab shows data rows",
    9:  "Next button on File Setup page is enabled and navigates to Config page",
    10: "Config page loads with Entity Level Setup section visible",
    11: "Entity Level Setup and Source Entities panel are both visible on Config page",
    12: "Save button is disabled before Data Owners and Data Stewards are selected",
    13: "Data Owner 'deepanraj.thangaraj@7dxperts.com' is selected and shown in the field",
    14: "Data Steward 'shreyas.senthilkumar@7dxperts.com' is selected and shown in the field",
    15: "Save button is enabled after both dropdowns are filled — click succeeds",
    16: "Success message 'Config updated successfully' is displayed after Save",
    17: "Next button is enabled after save and click navigates to next step",
    18: "Ingest page loads successfully after clicking Next on Config page",
    19: "'Create New' radio button is selected in Add to Data Domain Group section",
    20: "Save button is disabled before mandatory Ingest fields are filled",
    21: "Domain Name field accepts the dynamic value (filename without extension)",
    22: "Workflow Name field accepts the same value as Domain Name",
    23: "Workflow Type dropdown shows 'Monthly' as the selected value",
    24: "Expected Runtime (M) field accepts value '5'",
    25: "Notify on Delay checkbox is checked",
    26: "Save button is enabled after all fields are filled — click succeeds",
    27: "Success message 'Schedule data is saved successfully' is displayed",
    28: "Deploy Pipeline button becomes enabled after successful save",
    29: "Deploy Pipeline button is clicked successfully",
    30: "Deployment success message 'Successfully deployed' is displayed",
}


def _step_table(doc, steps: list[dict]):
    """
    steps: list of dicts with keys:
      step, description, expected, actual, status, screenshot_path, error (optional)
    """
    headers = ["#", "Description", "Expected Result", "Actual Result", "Status"]
    col_widths = [Inches(0.35), Inches(2.1), Inches(1.8), Inches(1.8), Inches(0.55)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_bg(hdr_cells[i], "1F3864")
        _cell_text(hdr_cells[i], h, bold=True, color="FFFFFF", size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        hdr_cells[i].width = col_widths[i]

    # Data rows
    for step in steps:
        is_fail = step.get("status", "PASS") == "FAIL"
        row = table.add_row()
        vals = [
            str(step.get("step", "")),
            step.get("description", ""),
            step.get("expected", ""),
            step.get("actual", ""),
            "✅ PASS" if not is_fail else "❌ FAIL",
        ]
        bg = "FFE7E7" if is_fail else "FFFFFF"
        for i, val in enumerate(vals):
            _set_cell_bg(row.cells[i], bg)
            color = "C00000" if is_fail and i == 4 else ("375623" if i == 4 else None)
            bold = i == 4
            _cell_text(row.cells[i], val, bold=bold, color=color, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER if i in (0, 4) else WD_ALIGN_PARAGRAPH.LEFT)
            row.cells[i].width = col_widths[i]

    doc.add_paragraph()


# ── Network section ───────────────────────────────────────────────────────────

def _network_section(doc, network_data: dict):
    _heading(doc, "Network Request & Response (Failed Step)", level=1)

    if not network_data:
        doc.add_paragraph("No network data captured.")
        return

    summary = network_data.get("network_summary", [])
    doc.add_paragraph(
        f"Total requests captured: {network_data.get('total_requests', 0)} | "
        f"Responses: {network_data.get('total_responses', 0)}"
    )

    # Show last few relevant entries (API calls, not static assets)
    api_calls = [
        e for e in summary
        if any(k in e.get("url", "").lower() for k in ["/api", "/v1", "/v2", "graphql", "query", "upload"])
        or int(e.get("status", 200)) >= 400
    ] or summary[-5:]  # fallback to last 5

    for idx, entry in enumerate(api_calls[-10:], 1):
        status_code = int(entry.get("status", 200))
        is_error = status_code >= 400

        doc.add_heading(f"API Call {idx}: {entry.get('method', '')} {entry.get('url', '')[:80]}", level=2)

        rows = [
            ("URL", entry.get("url", "N/A")),
            ("Method", entry.get("method", "N/A")),
            ("HTTP Status", f"{status_code} {entry.get('statusText', '')}"),
            ("Response Time", entry.get("responseTime", "N/A")),
        ]
        if entry.get("requestPayload"):
            rows.append(("Request Payload", str(entry["requestPayload"])[:500]))

        _summary_table(doc, rows)

        # Request Headers (collapsed to key ones)
        req_headers = entry.get("requestHeaders", {})
        if req_headers:
            doc.add_paragraph("Request Headers:", style="Normal").bold = True
            header_table = doc.add_table(rows=1, cols=2)
            header_table.style = "Table Grid"
            _set_cell_bg(header_table.rows[0].cells[0], "D9E2F3")
            _cell_text(header_table.rows[0].cells[0], "Header", bold=True, size=8)
            _set_cell_bg(header_table.rows[0].cells[1], "D9E2F3")
            _cell_text(header_table.rows[0].cells[1], "Value", bold=True, size=8)
            for k, v in list(req_headers.items())[:10]:
                hr = header_table.add_row()
                _cell_text(hr.cells[0], k, size=8)
                _cell_text(hr.cells[1], str(v)[:200], size=8)
            doc.add_paragraph()

        # Response Headers
        resp_headers = entry.get("responseHeaders", {})
        if resp_headers:
            doc.add_paragraph("Response Headers:", style="Normal")
            resp_table = doc.add_table(rows=1, cols=2)
            resp_table.style = "Table Grid"
            _set_cell_bg(resp_table.rows[0].cells[0], "D9E2F3" if not is_error else "FFE7E7")
            _cell_text(resp_table.rows[0].cells[0], "Header", bold=True, size=8)
            _set_cell_bg(resp_table.rows[0].cells[1], "D9E2F3" if not is_error else "FFE7E7")
            _cell_text(resp_table.rows[0].cells[1], "Value", bold=True, size=8)
            for k, v in list(resp_headers.items())[:10]:
                hr = resp_table.add_row()
                _cell_text(hr.cells[0], k, size=8)
                _cell_text(hr.cells[1], str(v)[:200], size=8)
            doc.add_paragraph()

    # JSON file references
    req_file = network_data.get("request_file", "")
    resp_file = network_data.get("response_file", "")
    if req_file and resp_file:
        note = doc.add_paragraph(
            f"📁 Full JSON saved:\n  • {os.path.basename(req_file)}\n  • {os.path.basename(resp_file)}"
        )
        for run in note.runs:
            run.font.size = Pt(9)
            run.font.italic = True
    doc.add_paragraph()


# ── Main generator ─────────────────────────────────────────────────────────────

def generate_execution_report(
    test_name: str,
    test_id: str,
    passed: bool,
    steps: list[dict],
    screenshots: list[dict],
    run_dir: str,
    failed_step: int = 0,
    failed_step_description: str = "",
    error_message: str = "",
    network_data: dict = None,
    output_dir: str = "reports/execution_reports",
) -> str:
    """
    Generate a full Word execution report for PASS or FAIL scenarios.

    Args:
        test_name:               Test function name
        test_id:                 Human-readable TC ID (e.g., "TC-01")
        passed:                  True = PASS, False = FAIL
        steps:                   List of step result dicts
        screenshots:             List of {step, path, description, is_failure}
        run_dir:                 Folder where screenshots are stored
        failed_step:             Step number that failed (0 if passed)
        failed_step_description: Description of failed step
        error_message:           Exception/assertion text
        network_data:            Dict from NetworkCapture.capture_network_for_step()
        output_dir:              Where to save the .docx

    Returns:
        Absolute path to generated .docx
    """
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    timestamp_file = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    safe_test_id = test_id.replace("-", "")
    if passed:
        filename = f"{safe_test_id}_Evidence_{timestamp_file}.docx"
    else:
        filename = f"{safe_test_id}_FAILED_{timestamp_file}.docx"
        
    output_path = os.path.join(output_dir, filename)

    doc = Document()

    # Page margins (narrower for more content space)
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Title ──────────────────────────────────────────────────────────────────
    title = doc.add_heading("Test Execution Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_paragraph("DataParQ Automation Framework")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        run.font.italic = True
    doc.add_paragraph()

    # ── Status Banner ──────────────────────────────────────────────────────────
    _status_banner(doc, passed)

    # ── Summary Table ──────────────────────────────────────────────────────────
    _heading(doc, "Test Case Summary", level=1)
    summary_rows = [
        ("Test Case ID",    test_id),
        ("Test Case Name",  test_name),
        ("Execution Date",  timestamp),
        ("Overall Status",  "✅ PASSED" if passed else "❌ FAILED"),
        ("Total Steps",     str(len(steps))),
        ("Steps Passed",    str(sum(1 for s in steps if s.get("status") == "PASS"))),
        ("Steps Failed",    str(sum(1 for s in steps if s.get("status") == "FAIL"))),
    ]
    if not passed:
        summary_rows += [
            ("Failed Step",         f"Step {failed_step}"),
            ("Failure Reason",      error_message[:300] if error_message else "Unknown"),
        ]
    _summary_table(doc, summary_rows)

    # ── Step-by-Step Table ────────────────────────────────────────────────────
    _heading(doc, "Step-by-Step Execution Details", level=1)
    _step_table(doc, steps)

    # ── Screenshots per step ──────────────────────────────────────────────────
    _heading(doc, "Step Screenshots", level=1)

    screenshot_map = {s["step"]: s for s in screenshots}

    for step in steps:
        snum = step.get("step", 0)
        sdesc = step.get("description", "")
        sstatus = step.get("status", "PASS")
        shot = screenshot_map.get(snum)

        doc.add_heading(
            f"Step {snum:02d} — {'✅' if sstatus == 'PASS' else '❌'} {sdesc}",
            level=2
        )

        if shot:
            label = f"{'⚠ FAILURE at ' if shot.get('is_failure') else ''}Step {snum}: {sdesc}"
            _add_image(doc, shot["path"], width_inches=5.8, caption=label)
        else:
            doc.add_paragraph("No screenshot available for this step.")

        doc.add_paragraph()

    # ── Network Section (FAIL only) ────────────────────────────────────────────
    if not passed and network_data:
        doc.add_page_break()
        _network_section(doc, network_data)

    # ── Footer ─────────────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph(
        f"Generated by DataParQ Automation Framework | {timestamp}"
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.italic = True

    doc.save(output_path)
    logger.info(f"INFO - Execution report saved: {output_path}")
    return output_path
